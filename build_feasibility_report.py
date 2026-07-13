from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from datetime import date
from pathlib import Path


OUT = Path(r"D:\work\quest\モンスター育成RPG_PWA実現可能性検討書.docx")

NAVY = "183B56"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2B6F77"
INK = "263238"
MUTED = "66737B"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F3ED"
PALE_GOLD = "FFF5D9"
PALE_RED = "FBEAEA"
WHITE = "FFFFFF"
BORDER = "CBD4DA"
FONT = "Yu Gothic"


def set_run_font(run, size=None, bold=None, color=INK, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            w = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(w))
            cell.width = Inches(w / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_table_header(row)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    set_run_font(r, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    rr = paragraph.add_run()
    rr._r.append(fld_char1)
    rr._r.append(instr_text)
    rr._r.append(fld_char2)
    set_run_font(rr, size=9, color=MUTED)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_body(doc, text, bold_lead=None, after=6, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.keep_together = keep
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    set_run_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    set_run_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), label_color)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=label_color)
    r = p.add_run(text)
    set_run_font(r, color=INK)


def add_table(doc, headers, rows, widths, header_fill=PALE_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(text))
            set_run_font(r, size=font_size, color=INK)
            if i == 0 and len(headers) <= 3:
                r.bold = True
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

# standard_business_brief preset with named Japanese-font override (Yu Gothic).
styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for level, size, color, before, after in [
    (1, 16, BLUE, 16, 8),
    (2, 13, BLUE, 12, 6),
    (3, 12, DARK_BLUE, 8, 4),
]:
    st = styles[f"Heading {level}"]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name, left, first in [("List Bullet", 0.5, -0.25), ("List Bullet 2", 0.8, -0.2), ("List Number", 0.5, -0.25)]:
    st = styles[name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(left)
    st.paragraph_format.first_line_indent = Inches(first)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.167

# Running header/footer (quiet and omitted on first page).
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
set_run_font(hp.add_run("MONSTER-BREEDING RPG / PWA FEASIBILITY"), size=8.5, bold=True, color=MUTED)
footer = section.footer
add_page_number(footer.paragraphs[0])
first_footer = section.first_page_footer
add_page_number(first_footer.paragraphs[0])

# Editorial cover.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(54)
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("FEASIBILITY REPORT"), size=10, bold=True, color=TEAL)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("モンスター育成RPG"), size=30, bold=True, color=NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("iPhone向けPWA 実現可能性検討書"), size=18, color=DARK_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
set_run_font(p.add_run("Macを使わず、個人で遊べる作品を現実的な規模で完成させるための判断材料"), size=11, color=MUTED)

add_callout(doc, "結論", "個人用・1人プレイ・2D・ターン制の範囲なら実現可能性は高い。まずはPWAで縦切り版を作り、既存作品の素材や固有名詞は使わない。", fill=PALE_GREEN, label_color=TEAL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after = Pt(4)
set_run_font(p.add_run("作成日  2026年7月12日"), size=10, color=MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("前提：販売なし／個人利用／Windowsで開発／iPhoneでプレイ"), size=9.5, color=MUTED)

doc.add_page_break()

add_heading(doc, "1. エグゼクティブサマリー", 1)
add_body(doc, "本企画は、モンスターを仲間にし、育成し、能力を継承させながらターン制戦闘を進める、オリジナルのモンスター育成RPGとして定義すれば、MacなしでもPWAとして十分に制作できる。iPhoneのホーム画面から独立したWebアプリとして起動でき、オフライン動作、ローカルセーブ、タッチ操作、BGM・効果音も実装可能である。")
add_body(doc, "一方、ネイティブiOSアプリと同じバックグラウンド常駐、完全に消えない端末内データ、App Store配布、Xcodeによる詳細デバッグは期待しない方がよい。ゲーム設計をPWAの特性に合わせれば、これらは致命的な制約ではない。")

add_table(doc,
    ["判断項目", "評価", "要点"],
    [
        ["技術的実現性", "高い", "2D・1人用・ターン制はWeb技術と相性がよい"],
        ["Macなし開発", "可能", "Windowsで開発・公開し、iPhone実機で確認する"],
        ["完成可能性", "中〜高", "MVPを絞れば現実的。広いフィールドや通信は後回し"],
        ["権利リスク", "要対策", "固有名詞・画像・音楽・データの流用を避け、オリジナル化する"],
        ["推奨開始案", "Go", "React中心の静的PWAで縦切り版を先に作る"],
    ], [2100, 1200, 6060], font_size=9.5)

add_callout(doc, "推奨方針", "『ドラゴンクエストモンスターズを再現する』ではなく、『仲間化・育成・能力継承の楽しさを持つオリジナルRPGを作る』と定義する。", fill=PALE_GOLD, label_color="7A5A00")

add_heading(doc, "2. 企画の境界", 1)
add_heading(doc, "2.1 最初に作るもの", 2)
for item in [
    "縦画面を基本とする、1人用の2Dターン制RPG",
    "探索地点を選ぶ → 戦闘 → 仲間化 → 育成 → 継承合成 → 次の地点、という短い循環",
    "20〜30体のオリジナルモンスター、10〜15個のスキル、3体程度のボス",
    "端末内セーブ、JSON形式のバックアップ書き出し・読み込み",
    "ホーム画面起動と、一度読み込んだ後のオフラインプレイ",
]:
    add_bullet(doc, item)

add_heading(doc, "2.2 最初は作らないもの", 2)
for item in [
    "オンライン対戦、ランキング、チャット、フレンド機能",
    "ユーザー登録、課金、ガチャ、運営用管理画面",
    "広大な自由移動マップ、リアルタイム戦闘、3D表現",
    "iCloud同期、プッシュ通知、App Store公開",
    "既存作品のモンスター画像、名称、BGM、効果音、文章、数値データの転用",
]:
    add_bullet(doc, item)

add_body(doc, "この切り分けは機能削減ではなく、最も価値の高い遊びを先に検証するための順序付けである。縦切り版が面白く、iPhoneで安定して動いてから、フィールド探索やクラウドセーブを足す。")

add_heading(doc, "3. iPhone向けPWAとしての実現性", 1)
add_body(doc, "AppleのWebKitは、ホーム画面に追加したWebサイトをアプリのように独立起動する仕組みを提供している。iOS 26では、ホーム画面追加時に原則としてWebアプリとして開く挙動が案内されている。したがって、個人利用のゲームを『ホーム画面のアイコンから起動する』要件は満たせる。［1］")

add_table(doc,
    ["要件", "可否", "設計上の扱い", "注意点"],
    [
        ["ホーム画面起動", "○", "manifestとアイコンを用意", "追加操作は利用者が行う"],
        ["オフライン動作", "○", "Service Workerで本体と素材をキャッシュ", "初回読込はオンライン"],
        ["ターン制戦闘", "○", "DOM/CSSまたはCanvasで実装", "重い3Dは対象外"],
        ["端末内セーブ", "○", "IndexedDBに状態を保存", "削除・退避に備えバックアップ必須"],
        ["BGM・効果音", "○", "初回タップ後に音声を開始", "復帰時にAudioContextを再開"],
        ["バックグラウンド進行", "△", "経過時刻から復帰時に再計算", "常時実行を前提にしない"],
        ["プッシュ通知", "○", "必要なら後日追加", "ホーム画面Webアプリ・許可操作が必要［2］"],
        ["確実な永続保存", "△", "persist要求＋外部バックアップ", "Webデータは退避され得る［3］"],
    ], [1650, 750, 3400, 3560], font_size=8.8)

add_heading(doc, "3.1 PWA固有の重要対策", 2)
add_number(doc, "ゲーム内の意味のある操作ごとに自動保存する。戦闘開始・ターン確定・勝敗確定・仲間化・合成の各境界で保存する。")
add_number(doc, "起動時にセーブのスキーマ版を確認し、古いデータを移行する。更新によってセーブが読めなくなる事故を防ぐ。")
add_number(doc, "navigator.storage.persist() を要求し、StorageManager.estimate() で使用量を診断表示する。ただし永続化を保証とはみなさない。［3］")
add_number(doc, "『バックアップを書き出す』『バックアップから復元』を設定画面に常設する。ファイル名には日時とセーブ版を含める。")
add_number(doc, "visibilitychange と pageshow を監視し、画面復帰時に音声・表示・未完了トランザクションを再同期する。iOSでは背景移行時にWeb Audioが停止する挙動がある。［5］")

add_callout(doc, "設計原則", "バックグラウンドで動き続けるゲームではなく、いつ中断されても同じ状態から再開できるゲームにする。", fill=PALE_BLUE)

add_heading(doc, "4. 権利面の整理", 1)
add_body(doc, "販売せず自分だけで遊ぶ場合でも、『非営利だから自由に使える』とは限らない。日本の著作権法は私的使用のための複製に例外を設ける一方、インターネットへのアップロードは公衆送信・送信可能化に関係し、二次創作は複製権・翻案権等の許諾が原則として必要と文化庁が説明している。［6］［7］")
add_body(doc, "PWAは通常HTTPSのサーバーに配置するため、URLを広く告知しなくても、他人の画像・音楽等をサーバーに置く設計は避けるのが安全である。認証を付ければリスクがゼロになる、という扱いもしない。本節は一般的な企画判断であり、個別案件の法律相談ではない。")

add_table(doc,
    ["要素", "推奨", "理由"],
    [
        ["作品タイトル", "完全に別名", "公式作品・系列作品と誤認される表示を避ける"],
        ["モンスター名・外見", "すべてオリジナル", "キャラクター表現の複製・翻案リスクを避ける"],
        ["画像・アニメーション", "自作または利用許諾済み素材", "ゲーム画面の切り抜き・トレースは使わない"],
        ["BGM・効果音", "自作、CC0、または明確なゲーム利用ライセンス", "原曲・耳コピ・抽出音源は使わない"],
        ["システム", "抽象的な着想を再設計", "固有の文章、画面構成、配合表、数値の丸写しを避ける"],
        ["説明文", "自分の言葉で新規作成", "図鑑説明・スキル文言の転用を避ける"],
    ], [1900, 2900, 4560], font_size=9.0)

add_callout(doc, "赤線", "『ドラゴンクエスト』『ドラゴンクエストモンスターズ』の名称、公式ロゴ、既存モンスター、公式画像、BGM・効果音、図鑑文、配合表を成果物に入れない。", fill=PALE_RED, label_color="9B1C1C")
add_body(doc, "ゲームの着想として『モンスターを集め、育て、親の特徴を次世代へ継承する』という抽象的な遊びを研究し、名称、ルール、計算式、画面、文章、世界観を独自に設計する方針が、完成後も安心して保管・改良できる。")

add_heading(doc, "5. 推奨するゲーム仕様（MVP）", 1)
add_heading(doc, "5.1 コアループ", 2)
add_body(doc, "拠点で編成 → 探索地点を選択 → 1〜3戦のターン制バトル → 仲間化判定 → 成長・報酬 → 2体から新個体を誕生させ、特徴を継承 → 次の探索地点へ、という10〜15分単位の循環とする。自由移動マップは後回しにし、最初は地点選択式でゲーム性を検証する。")

add_heading(doc, "5.2 最小コンテンツ量", 2)
for item in [
    "モンスター24体：4系統 × 6体。各体に基本能力、特性1個、習得候補2〜3個を持たせる。",
    "スキル12個：攻撃4、回復2、補助4、状態異常2。最初から例外的な効果を増やしすぎない。",
    "探索8地点：各地点に通常敵、希少敵、報酬、解放条件をデータとして定義する。",
    "ボス3体：単なる高HPではなく、状態変化や攻撃予告で戦術を作る。",
    "継承合成：親2体から系統、能力補正、特性候補、スキル候補を決定。結果予告を表示し、理不尽な消失を避ける。",
]:
    add_bullet(doc, item)

add_heading(doc, "5.3 完成条件", 2)
for item in [
    "新規開始から最初のボス撃破まで、説明なしで15〜30分遊べる",
    "アプリを閉じ、機内モードで再起動しても続きから遊べる",
    "戦闘中の任意のタイミングで中断しても、不正な増殖や進行不能が起きない",
    "バックアップを書き出し、セーブを削除した後に復元できる",
    "iPhoneの小さい画面で、主要操作のタップ領域が十分に大きい",
    "既存作品の素材・名称・文章が成果物に含まれていない",
]:
    add_bullet(doc, item)

add_heading(doc, "6. 推奨技術構成", 1)
add_body(doc, "最初からゲームエンジンとUIフレームワークを二重に抱えると、画面遷移、状態同期、タッチ入力の責任範囲が複雑になる。MVPはReact + TypeScriptを中心に、HTML/CSSでメニューと戦闘画面を作る。自由移動マップが本当に必要になった時だけ、Phaser等のCanvasエンジンを追加する。")

add_table(doc,
    ["層", "推奨", "役割"],
    [
        ["画面", "React + TypeScript", "編成、戦闘、図鑑、合成、設定。縦画面・タッチ優先"],
        ["ビルド", "Vite", "Windowsで高速な開発サーバーと静的ファイル生成"],
        ["ゲームロジック", "純粋なTypeScript", "戦闘計算、仲間化、成長、合成。UIから分離してテスト"],
        ["マスターデータ", "JSON / TypeScript", "モンスター、技、敵編成、報酬。コード変更なしで調整"],
        ["セーブ", "IndexedDB", "セーブ本体、設定、マイグレーション情報"],
        ["オフライン", "Service Worker", "アプリ本体・画像・音声をバージョン単位でキャッシュ"],
        ["品質確認", "Vitest + Playwright", "計算の単体テストと主要画面の回帰確認"],
        ["公開", "静的HTTPSホスティング", "サーバーロジックなし。URLからインストール"],
    ], [1800, 2400, 5160], font_size=8.9)

add_heading(doc, "6.1 データ設計で先に決めること", 2)
for item in [
    "乱数シード：戦闘の再開と不正な引き直しをどう扱うか",
    "セーブの版番号：schemaVersion、appVersion、savedAtを必ず保存",
    "IDの安定性：表示名ではなく変更されない内部IDで参照",
    "計算の再現性：UIから直接乱数や能力値を書き換えない",
    "更新戦略：新しいService Workerを戦闘中に即時適用しない",
]:
    add_bullet(doc, item)

add_heading(doc, "6.2 構成イメージ", 2)
add_body(doc, "画面コンポーネントはゲーム状態を表示し、ユーザーのコマンドをゲームロジックへ渡す。ロジックは次の状態とイベント履歴を返し、セーブ層がIndexedDBへ原子的に保存する。Service Workerは実行コードと素材のオフライン供給だけを担当し、セーブ本体はキャッシュに入れない。")
add_callout(doc, "責務分離", "UI → コマンド → ゲームロジック → 次状態 → IndexedDB。表示と計算を分けると、バランス調整とバグ修正が大幅に楽になる。", fill=PALE_GREEN, label_color=TEAL)

add_heading(doc, "7. Macなしでの開発・テスト", 1)
add_body(doc, "ソースコード作成、ローカル開発、テスト、ビルド、HTTPSホスティングへの公開はWindowsだけで行える。Apple Developer Programへの加入やXcodeも、PWAのホーム画面利用には不要である。Web PushについてもAppleはDeveloper Program加入不要と案内しているが、今回のMVPでは通知自体を採用しない。［2］")

add_heading(doc, "7.1 推奨する確認サイクル", 2)
add_number(doc, "Windows上のChrome/Edgeで、ゲームロジック、レスポンシブ画面、キーボード操作を高速に確認する。")
add_number(doc, "変更をプレビュー用HTTPS URLへ自動公開する。公開物は必ずオリジナル素材だけにする。")
add_number(doc, "実機iPhoneのSafariでURLを開き、ホーム画面へ追加して、縦横回転、ノッチ、安全領域、タップ、音声、復帰、機内モードを確認する。")
add_number(doc, "正式なiOS Web InspectorはSafariの開発メニューからiPhoneを選ぶ手順であり、Macがない環境では使いにくい。画面内デバッグログ、エラー書き出し、端末情報表示を開発版に用意する。［4］")

add_heading(doc, "7.2 実機テスト項目", 2)
for item in [
    "初回アクセス → ホーム画面追加 → standalone起動",
    "初回タップ後のBGM開始、消音設定、アプリ切替後の音声復帰",
    "機内モード起動、キャッシュ更新直後の再起動",
    "戦闘中断、端末ロック、メモリ不足を想定した再読込",
    "セーブの自動保存、バックアップ書き出し、復元、旧版データ移行",
    "safe-area-inset、文字拡大、ダークモード、低電力モード",
]:
    add_bullet(doc, item)

add_heading(doc, "8. 工数とロードマップ", 1)
add_body(doc, "個人開発の工数は、Web開発経験と素材制作の方法で大きく変わる。以下は既存素材の無断流用をせず、2Dの簡素なオリジナル素材を用いる場合の目安である。未経験者は1.5〜2倍程度を見込む。")

add_table(doc,
    ["段階", "主な成果", "累計目安", "出口条件"],
    [
        ["0. 企画固定", "名称、コアループ、赤線、データ項目", "4〜8時間", "1ページ仕様が決まる"],
        ["1. 技術スパイク", "PWA起動、1戦、IndexedDB保存、音声", "20〜40時間", "iPhone機内モードで再戦可能"],
        ["2. 縦切り版", "仲間化、編成、成長、合成、1ボス", "80〜140時間", "15〜30分の循環が面白い"],
        ["3. 個人用MVP", "24体、8地点、3ボス、バックアップ", "180〜320時間", "最初から最後まで遊べる"],
        ["4. 拡張", "探索マップ、演出、クラウドセーブ等", "400時間〜", "必要性を再評価して選択"],
    ], [1650, 3440, 1400, 2870], font_size=8.7)

add_body(doc, "週6〜8時間なら、縦切り版まで約3〜5か月、個人用MVPまで約6〜12か月が現実的な範囲となる。素材を凝りすぎるとコードより制作時間が増えるため、最初はシルエット、簡単な2D立ち絵、短いアニメーションでよい。")

add_heading(doc, "8.1 最初の4週間", 2)
add_number(doc, "第1週：作品名、世界観1段落、6体のモンスター、4スキル、戦闘式を紙上で決める。")
add_number(doc, "第2週：PWAの雛形、ホーム画面起動、1画面の戦闘、音声ON/OFFを作る。")
add_number(doc, "第3週：戦闘ロジックをUIから分離し、勝敗、状態異常、乱数を自動テストする。")
add_number(doc, "第4週：IndexedDB保存、機内モード、バックアップ書き出しを実機で通す。続行判断はここで行う。")

add_heading(doc, "9. 主なリスクと軽減策", 1)
add_heading(doc, "9.1 スコープ膨張", 2)
add_body(doc, "最も大きいリスクは技術ではなく、作りたい要素が増え続けること。『最初のボスまで』を独立した完成品として扱い、追加機能はその後の別マイルストーンにする。")
add_heading(doc, "9.2 セーブ消失", 2)
add_body(doc, "端末内データだけを唯一の正本にしない。永続化要求、自動保存、世代付きバックアップ、手動エクスポートを組み合わせる。クラウド同期はMVP後でよい。")
add_heading(doc, "9.3 iOSだけの不具合", 2)
add_body(doc, "Windowsブラウザだけで完成判定をしない。毎週少なくとも1回、ホーム画面版のiPhone実機で、音声、復帰、オフライン、セーブを通す。")
add_heading(doc, "9.4 権利侵害・公開事故", 2)
add_body(doc, "リポジトリと公開URLに、公式作品から抽出・転記した素材を一度も入れない。試作段階からオリジナルまたは明示的に許諾された素材だけを使う。")
add_heading(doc, "9.5 バランス調整の迷走", 2)
add_body(doc, "戦闘結果のイベントログと簡易シミュレーターを用意し、感覚だけで数値を変更しない。勝率、平均ターン数、使用スキル比率を確認できるようにする。")

add_heading(doc, "10. 最終提案", 1)
add_callout(doc, "Go / 条件付き", "PWA方式で開始してよい。条件は、オリジナルIP、1人用、地点選択式、24体程度、端末内セーブ＋バックアップを最初の完成線とすること。", fill=PALE_GREEN, label_color=TEAL)
add_body(doc, "Macがないことは本企画の中止理由にならない。むしろPWAに限定することで、App Store審査、署名、Xcode、年会費、ネイティブ更新の負担を避けられる。制約は、ゲームを中断耐性の高いターン制にし、セーブを防御的に扱い、実機確認を習慣化することで吸収できる。")

add_heading(doc, "10.1 次に決める5項目", 2)
for item in [
    "仮タイトルと世界観：既存作品と明確に異なるもの",
    "戦闘人数：1対1、3対3など。MVPは最大3対3まで",
    "仲間化方式：戦闘後確率、会話、アイテム等",
    "継承合成の方針：親を失うか、能力・特性を何個継ぐか",
    "グラフィック方針：簡素な2D、ピクセルアート、シルエット等",
]:
    add_bullet(doc, item)

add_body(doc, "この5項目が決まれば、次の作業は『技術スパイク』である。6体・4スキル・1戦・1セーブだけの最小プロジェクトを作り、iPhoneのホーム画面版で動かす。ここで得た実機結果を基に、正式な仕様書とタスク分解へ進む。")

add_heading(doc, "参考資料", 1)
sources = [
    ("［1］WebKit, WebKit Features in Safari 26.0（iOS 26のホーム画面Webアプリ挙動）", "https://webkit.org/blog/17333/webkit-features-in-safari-26-0/"),
    ("［2］WebKit, Web Push for Web Apps on iOS and iPadOS（ホーム画面Webアプリ、通知、Developer Program不要）", "https://webkit.org/blog/13878/web-push-for-web-apps-on-ios-and-ipados/"),
    ("［3］WebKit, Updates to Storage Policy（Quota、eviction、persistent mode）", "https://webkit.org/blog/14403/updates-to-storage-policy/"),
    ("［4］Apple Developer, Inspecting iOS and iPadOS（Home Screen Web Appsの検査手順）", "https://developer.apple.com/documentation/safari-developer-tools/inspecting-ios"),
    ("［5］WebKit Bugzilla 220879（iOSでのWeb Audio背景停止に関する実装記録）", "https://bugs.webkit.org/show_bug.cgi?id=220879"),
    ("［6］e-Gov法令検索, 著作権法（昭和四十五年法律第四十八号）", "https://laws.e-gov.go.jp/law/345AC0000000048"),
    ("［7］文化庁, ここが知りたい著作権（私的使用、公衆送信、二次創作の説明）", "https://www.bunka.go.jp/seisaku/chosakuken/taisetsu/point/index.html"),
]
for title, url in sources:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_hyperlink(p, title, url)

add_body(doc, "注：Webプラットフォームの仕様とiOSの挙動は更新される。実装開始時とiOS更新時には、対象端末で再確認すること。権利面の記述は一般的なリスク整理であり、法的助言ではない。", after=0)

# Core properties.
doc.core_properties.title = "モンスター育成RPG iPhone向けPWA 実現可能性検討書"
doc.core_properties.subject = "Macなしで制作する個人用モンスター育成RPGの技術・権利・工数評価"
doc.core_properties.author = ""
doc.core_properties.keywords = "PWA, iPhone, モンスター育成RPG, 実現可能性"

# Avoid orphan headings and set keep-with-next for table-adjacent labels globally.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True

doc.save(OUT)
print(OUT)
